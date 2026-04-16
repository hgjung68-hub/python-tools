import os
import time
import re
import numpy as np
import pandas as pd
import pdfplumber
import easyocr
import meilisearch
from docx import Document
from pyhwpx import Hwp
from datetime import datetime

# 1. 엔진 초기화
print("\n" + "="*60)
print("🚀 [CISS] Contract Integrated Search System 가동(단일폴더용)")
print("🚀 엔진 초기화 중 (EasyOCR, Meilisearch)...")
print("지원문서: .pdf, .doc, .docx, .xls, .xlsx, .hwp, .hwpx")
print("="*60)

# GPU가 없다면 False 유지, 있다면 True로 변경 시 속도 향상
reader = easyocr.Reader(['ko', 'en'], gpu=False)
# 1. Local PC 구동(26.2.27 생성) Admin API Key
# ADMIN_API_KEY = '095edf6d4d86a911ef7a33828f8369a86e962f528c9f1724ebc2729a85dda28b'
# 2. docker 구동(26.3.18 생성) Admin API Key
ADMIN_API_KEY = '66438e50779a8d2792f1f7e6c7ed70539dffb809d9e84fe09a000f1f8b7bbfb2' 
client = meilisearch.Client('http://127.0.0.1:7700', ADMIN_API_KEY)
index = client.index('ciss_contracts')

# 파싱대상 계약서 파일 폴더
base_path = "./Contract_3"

# --- [모듈 0] Excel 추출 (pandas 활용) ---
def extract_excel_safe(file_path):
    """
    엑셀의 모든 시트를 순회하며 텍스트 데이터를 추출
    """
    combined_text = []
    try:
        # 모든 시트 읽기 (engine='openpyxl' 권장)
        all_sheets = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
        for sheet_name, df in all_sheets.items():
            # 결측값 제거 및 문자열 변환
            clean_df = df.fillna("")
            sheet_content = clean_df.astype(str).values.flatten()
            combined_text.append(f"[{sheet_name} 시트]\n" + " ".join(sheet_content))
    except Exception as e:
        print(f"\n      ❌ Excel 에러: {e}")
    return "\n\n".join(combined_text)

# --- [모듈 1] PDF 추출 (상세 페이지 로그 유지) ---
def extract_pdf_optimized(file_path):
    full_text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_start = time.time()
                # 현재 페이지와 전체 페이지 표시
                print(f"      📄 [{i+1}/{total_pages}] 페이지 분석 중...", end=" ", flush=True)
                
                page_text = page.extract_text() or ""
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        page_text += "\n" + " ".join([str(c) for c in row if c is not None])
                
                # 조건부 OCR 로그 출력
                if len(page_text.strip()) < 50:
                    print("(이미지 감지 👉 OCR 실행)", end=" ", flush=True)
                    img = page.to_image(resolution=200).original
                    ocr_result = reader.readtext(np.array(img), detail=0, paragraph=True)
                    page_text += "\n" + " ".join(ocr_result)
                else:
                    print("(텍스트 스캔 완료)", end=" ", flush=True)

                full_text += page_text + "\n"
                print(f"[{time.time() - page_start:.1f}초]")
    except Exception as e:
        print(f"\n      ❌ PDF 에러: {e}")
    return full_text

# --- [모듈 2] Word 추출 (문단/이미지 상세 로그) ---
def extract_word_advanced(file_path):
    combined_content = []
    try:
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        print(f"      📝 본문 {len(paragraphs)}문단 분석 중...", end=" ", flush=True)
        for para in paragraphs:
            if para.text.strip(): combined_content.append(para.text)
        print("완료!")

        if doc.tables:
            print(f"      📊 표 {len(doc.tables)}개 추출 중...", end=" ", flush=True)
            for table in doc.tables:
                for row in table.rows:
                    combined_content.append(" ".join([cell.text for cell in row.cells if cell.text.strip()]))
            print("완료!")

        img_rels = [rel for rel in doc.part.rels.values() if "image" in rel.target_ref]
        if img_rels:
            print(f"      🖼️ 이미지 OCR 분석 ({len(img_rels)}개)...", end=" ", flush=True)
            for rel in img_rels:
                try:
                    ocr_result = reader.readtext(rel.target_part.blob, detail=0, paragraph=True)
                    if ocr_result: combined_content.append(" ".join(ocr_result))
                except: pass
            print("완료!")
    except Exception as e:
        print(f"\n      ❌ Word 에러: {e}")
    return "\n".join(combined_content)

# --- [모듈 3] HWP/HWPX 추출 (pyhwpx 활용) ---
def extract_hwp_advanced(file_path):
    """
    한글(hwp, hwpx) 문서의 본문 및 표 데이터를 추출
    (주의: PC에 한글 프로그램이 설치되어 있어야 제어 가능)
    """
    combined_text = []
    try:
        hwp = Hwp()
        hwp.open(file_path)
        
        # 1. 본문 전체 텍스트 추출
        full_text = hwp.get_text()
        if full_text.strip():
            combined_text.append("[본문 내용]\n" + full_text)
            
        # 2. 문서 내 표(Table) 데이터 정밀 추출
        tables = hwp.get_tables()
        if tables:
            combined_text.append(f"\n[표 데이터 ({len(tables)}개)]")
            for i, table in enumerate(tables):
                table_data = hwp.get_table_data(table)
                # 리스트 형태의 표 데이터를 문자열로 변환
                for row in table_data:
                    row_cells = []
                    for cell in row:
                        if cell is None: continue
                        
                        # 튜플 에러 해결 핵심 로직: 데이터 타입 체크
                        if isinstance(cell, (tuple, list)):
                            # 튜플인 경우 첫 번째 요소(텍스트)만 가져옴
                            cell_val = str(cell[0]).strip() if cell else ""
                        else:
                            # 일반 문자열이나 숫자인 경우
                            cell_val = str(cell).strip()
                        
                        if cell_val: row_cells.append(cell_val)
                    
                    if row_cells:
                        combined_text.append(" | ".join(row_cells))
        
        hwp.quit()
    except Exception as e:
        # 한글 프로그램 미설치 또는 OLE 오류 대비용 fallback
        print(f"\n      ❌ HWP 제어 에러: {e} (표준 텍스트 스캔 시도)")
        try:
            # pyhwpx 실패 시 최소한의 텍스트라도 건지기 위한 로직
            import olefile
            f = olefile.OleFileIO(file_path)
            if f.exists('PrvText'):
                raw = f.openstream('PrvText').read().decode('utf-16')
                combined_text.append("[미리보기 텍스트 추출]\n" + raw)
        except:
            return f"[HWP 파싱 실패: {os.path.basename(file_path)}]"
    finally:
        # 에러 여부와 관계없이 프로세스를 종료하여 파일 점유 해제
        if hwp is not None:
            try:
                hwp.quit()
                # 윈도우가 파일 잠금을 완전히 풀 때까지 아주 짧은 대기시간
                time.sleep(0.5) 
            except:
                pass
            
    return "\n".join(combined_text)

# --- [메인 배치 로직: 폴더 스캔 및 이름 변경] ---
def run_ciss_batch_processing():
    # base_path = "./Contract"
    if not os.path.exists(base_path):
        print(f"❌ 경로 오류: {base_path} 폴더가 없습니다.")
        return

    # 미처리 폴더 스캔
    folders = [f for f in os.listdir(base_path) 
               if os.path.isdir(os.path.join(base_path, f)) 
               and not f.startswith(('Ok-', 'Err-'))]

    total_folders = len(folders)
    print(f"\n📂 총 {total_folders}개의 신규 계약 폴더를 처리합니다.")

    for idx, folder_name in enumerate(folders):
        folder_start = time.time()
        contract_no = folder_name
        folder_full_path = os.path.join(base_path, folder_name)
        
        print(f"\n[{idx+1}/{total_folders}] 📂 계약번호: {contract_no}")
        print(f"  {'='*50}")
        
        all_contents = []
        process_success = True
        
        try:
            files = [f for f in os.listdir(folder_full_path) if os.path.isfile(os.path.join(folder_full_path, f))]
            if not files:
                raise Exception("폴더 내에 파일이 없습니다.")

            for f_idx, f_name in enumerate(files):
                f_path = os.path.join(folder_full_path, f_name)
                ext = os.path.splitext(f_name)[1].lower()
                
                print(f"  📁 ({f_idx+1}/{len(files)}) 파일: {f_name}")

                # 파일별 추출 시도
                try:
                    if ext == '.pdf':
                        all_contents.append(extract_pdf_optimized(f_path))
                    elif ext in ['.docx', '.doc']:
                        all_contents.append(extract_word_advanced(f_path))
                    elif ext in ['.xlsx', '.xls']:
                        print(f"      📊 Excel 데이터 추출 중...", end=" ", flush=True)
                        excel_text = extract_excel_safe(f_path)
                        all_contents.append(excel_text)
                        print("완료!")
                    elif ext in ['.hwp', '.hwpx']:
                        print(f"      📝 HWP 데이터 추출 중...", end=" ", flush=True)
                        all_contents.append(extract_hwp_advanced(f_path))
                        print("완료!")
                    else:
                        print(f"      ⏩ 지원하지 않는 확장자입니다.")
                except Exception as file_e:
                    print(f"      ⚠️ 파일 개별 처리 중 오류: {file_e}")
                    process_success = False # 하나라도 실패하면 최종 상태는 Err-

            # Meilisearch 적재 (하나라도 내용이 있으면 시도)
            if all_contents:
                print(f"  🚀 Meilisearch 데이터 전송...", end=" ", flush=True)
                doc_payload = {
                    "contract_no": contract_no,
                    "content_text": "\n\n".join(all_contents),
                    "file_list": files,
                    "indexed_at": datetime.now().isoformat()
                }
                index.add_documents([doc_payload], primary_key='contract_no')
                print("성공!")
            else:
                raise Exception("추출된 텍스트 내용이 전혀 없습니다.")
            
        except Exception as e:
            print(f"\n  ❌ 최종 에러: {e}")
            process_success = False

        # 폴더명 변경 및 결과 요약
        status_prefix = "Ok-" if process_success else "Err-"
        new_folder_path = os.path.join(base_path, f"{status_prefix}{contract_no}")
        
        try:
            # 중복 폴더 방지 처리
            if os.path.exists(new_folder_path):
                new_folder_path += f"_{int(time.time())}"
            os.rename(folder_full_path, new_folder_path)
            duration = time.time() - folder_start
            print(f"  ✨ {status_prefix}처리 완료 (소요시간: {duration:.1f}초)")
        except Exception as e:
            print(f"  ⚠️ 폴더명 변경 실패: {e}")
        print(f"  {'='*50}")

    print(f"\n🏁 모든 배치 작업이 끝났습니다.")

if __name__ == "__main__":
    run_ciss_batch_processing()